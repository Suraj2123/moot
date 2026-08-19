"""Turn an uploaded file into note text.

Everything downstream of this module -- chunking, embedding, retrieval, the
chat -- already works on plain text. So the whole job here is to get plain
text out of the handful of formats a student actually has lecture material
in, and to fail in a way that says what went wrong.

The failure messages matter more than they look. A PDF that yields nothing is
almost always a scan, and "no text found" sends someone off to check their
file for corruption when the real answer is "this is a picture of words and
OCR is not built yet". Each error below names the actual cause.

What is deliberately not here:

  * OCR. It needs a model or a service, it is slow, and it is wrong often
    enough that a silent bad transcript is worse than a refusal.
  * .doc, .pages, .rtf. Old Word is a different format from .docx entirely,
    and the others are rare enough to not be worth the dependency.
  * Anything that executes. No macros are run, no embedded content is
    followed. A document is bytes to be read, never instructions.
"""

from __future__ import annotations

import io
import logging
import re
import zipfile
from dataclasses import dataclass

logger = logging.getLogger(__name__)

# 10 MB. A semester of lecture notes is far under this; a video someone
# renamed is far over. The cap is enforced while reading, not after, so a
# huge upload cannot fill memory before being rejected.
MAX_BYTES = 10 * 1024 * 1024

# Extraction is bounded by page count too. A 900-page textbook is not a note,
# and parsing one blocks a request for minutes.
MAX_PDF_PAGES = 300

SUPPORTED = {
    ".md": "Markdown",
    ".markdown": "Markdown",
    ".txt": "Plain text",
    ".text": "Plain text",
    ".pdf": "PDF",
    ".docx": "Word",
}


class DocumentError(ValueError):
    """Something about this file stops it becoming a note. Safe to show."""


@dataclass(frozen=True)
class Extracted:
    text: str
    kind: str
    #: Set when the file parsed but something about the result is worth
    #: saying -- a truncated PDF, a mostly-empty document.
    notice: str | None = None


def suffix_of(filename: str) -> str:
    name = (filename or "").strip().lower()
    dot = name.rfind(".")
    return name[dot:] if dot > 0 else ""


def is_supported(filename: str) -> bool:
    return suffix_of(filename) in SUPPORTED


def title_from_filename(filename: str) -> str:
    """A readable default title, since most uploads will not carry one.

    `CS101 - week 3 lecture (final).pdf` should become
    `CS101 - week 3 lecture (final)`, not the raw filename with an extension
    a user then has to delete by hand.
    """
    name = (filename or "").strip().replace("\\", "/").rsplit("/", 1)[-1]
    suffix = suffix_of(name)
    if suffix:
        name = name[: -len(suffix)]
    name = name.replace("_", " ").strip()
    name = re.sub(r"\s+", " ", name)
    return name or "Untitled upload"


def extract(filename: str, data: bytes) -> Extracted:
    """Plain text from an uploaded file, or a DocumentError saying why not."""
    suffix = suffix_of(filename)
    if suffix not in SUPPORTED:
        supported = ", ".join(sorted(SUPPORTED))
        raise DocumentError(
            f"{suffix or 'That file type'} is not supported. Upload one of: {supported}."
        )
    if not data:
        raise DocumentError("That file is empty.")
    if len(data) > MAX_BYTES:
        raise DocumentError(
            f"That file is {_mb(len(data))}, over the {_mb(MAX_BYTES)} limit."
        )

    if suffix == ".pdf":
        result = _from_pdf(data)
    elif suffix == ".docx":
        result = _from_docx(data)
    else:
        result = _from_text(data, SUPPORTED[suffix])

    if not result.text.strip():
        raise DocumentError(
            f"No text could be read from that {SUPPORTED[suffix]} file. "
            "If it is a scan or photographed pages, the text has to be "
            "recognised first -- moot does not do OCR yet."
        )
    return result


# ------------------------------------------------------------------ formats


def _from_text(data: bytes, kind: str) -> Extracted:
    """Decode, trying the encodings that actually turn up.

    UTF-8 covers nearly everything. The fallbacks exist because a file
    exported from an older Windows tool is cp1252, and failing on one smart
    quote would be a silly reason to reject a whole set of notes.

    The binary check has to come first, because latin-1 decodes *any* byte
    sequence -- it maps all 256 values. Without this, a binary file renamed
    to .txt would not fail to decode, it would succeed into mojibake and be
    stored as a note. A refusal is much better than a note full of noise that
    then gets embedded and matched against.
    """
    if b"\x00" in data:
        # The heuristic git uses, and it is right for the same reason: text
        # essentially never contains NUL, and binary formats almost always do.
        raise DocumentError(
            "That file looks like binary data rather than text. If it is a "
            "document, save it as PDF, .docx, or .txt first."
        )

    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return Extracted(text=data.decode(encoding), kind=kind)
        except UnicodeDecodeError:
            continue

    # Unreachable while latin-1 is in the list above; kept so that removing it
    # fails loudly rather than falling off the end returning None.
    raise DocumentError("That file is not readable as text.")


def _from_pdf(data: bytes) -> Extracted:
    if not data.startswith(b"%PDF"):
        raise DocumentError(
            "That file is named .pdf but is not a PDF. Check you uploaded the "
            "right file."
        )
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocumentError("This server cannot read PDFs: pypdf is not installed.") from exc

    try:
        reader = PdfReader(io.BytesIO(data))
    except Exception as exc:  # noqa: BLE001 - pypdf raises several types
        raise DocumentError(
            "That PDF could not be opened. It may be damaged or incomplete."
        ) from exc

    if getattr(reader, "is_encrypted", False):
        # An empty-password PDF is common (print-protected, not secret) and
        # decrypts silently. A real password is a refusal, not a puzzle.
        try:
            if not reader.decrypt(""):
                raise DocumentError(
                    "That PDF is password protected. Remove the password and "
                    "upload it again."
                )
        except DocumentError:
            raise
        except Exception as exc:  # noqa: BLE001
            raise DocumentError("That PDF is password protected.") from exc

    pages = reader.pages
    notice = None
    if len(pages) > MAX_PDF_PAGES:
        notice = (
            f"Only the first {MAX_PDF_PAGES} of {len(pages)} pages were read."
        )
        pages = pages[:MAX_PDF_PAGES]

    parts: list[str] = []
    for number, page in enumerate(pages, start=1):
        try:
            parts.append(page.extract_text() or "")
        except Exception:  # noqa: BLE001 - one bad page should not lose the rest
            logger.warning("could not extract page %d of a PDF", number)
            parts.append("")

    return Extracted(text=_tidy("\n\n".join(parts)), kind="PDF", notice=notice)


def _from_docx(data: bytes) -> Extracted:
    if not data.startswith(b"PK"):
        # .docx is a zip. A .doc renamed to .docx is the usual cause, and it
        # is worth naming because the fix is one Save As.
        raise DocumentError(
            "That file is named .docx but is not a Word document. Older .doc "
            "files need to be saved as .docx first."
        )
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise DocumentError(
            "This server cannot read Word files: python-docx is not installed."
        ) from exc

    try:
        document = docx.Document(io.BytesIO(data))
    except (zipfile.BadZipFile, KeyError) as exc:
        raise DocumentError(
            "That Word document could not be opened. It may be damaged."
        ) from exc
    except Exception as exc:  # noqa: BLE001
        raise DocumentError("That Word document could not be read.") from exc

    parts = [p.text for p in document.paragraphs]

    # Tables hold real content in lecture handouts -- definitions, schedules,
    # comparisons -- and python-docx does not include them in `paragraphs`.
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))

    return Extracted(text=_tidy("\n\n".join(parts)), kind="Word")


# ------------------------------------------------------------------ tidying


def _tidy(text: str) -> str:
    """Normalise the whitespace extractors produce.

    PDF extraction in particular emits ragged spacing and long runs of blank
    lines, which survive into chunking and waste room in every chunk they land
    in. Paragraph breaks are kept because chunking is paragraph-aware.
    """
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r" *\n *", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def _mb(size: int) -> str:
    return f"{size / (1024 * 1024):.1f} MB"
